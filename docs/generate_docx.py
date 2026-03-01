#!/usr/bin/env python3
"""Generate NitNab Design Review .docx document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# -- Page Setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica Neue'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Helvetica Neue'
    hs.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
    if level == 1:
        hs.font.size = Pt(24)
    elif level == 2:
        hs.font.size = Pt(18)
    else:
        hs.font.size = Pt(14)

ACCENT = RGBColor(0x00, 0x71, 0xE3)  # Apple blue

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# =====================================================================
# COVER PAGE
# =====================================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NitNab')
run.font.size = Pt(42)
run.font.color.rgb = ACCENT
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Design Review & Feature Implementation Plan')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Prepared by: Design & Engineering Review Team\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)
run = meta.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)
run = meta.add_run('Version: 1.0 | Confidential')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Product Audit: Current State Assessment',
    '3. Competitive Landscape Analysis',
    '4. Platform Alignment: macOS 26 Tahoe & Liquid Glass',
    '5. Feature Plan: Three Pillars of Improvement',
    '6. UX Polish & Usability Recommendations',
    '7. Accessibility & Inclusivity',
    '8. Implementation Roadmap',
    '9. Appendix: Issue Registry & Sources',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE SUMMARY
# =====================================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'NitNab is a macOS 26+ speech transcription application built on Apple\'s native '
    'SpeechTranscriber API and Apple Intelligence Foundation Models. It offers dual-mode UI '
    '(Standard and Advanced), company/vocabulary context for improved accuracy, iCloud-first '
    'persistence, and AI-powered post-processing.'
)

doc.add_paragraph(
    'This review identifies 26 UX issues across the product, benchmarks NitNab against four '
    'leading competitors (Otter.ai, MacWhisper, Descript, Rev.com), and delivers a revised '
    'implementation plan for three user-requested features:'
)

features = [
    ('Per-File Transcription', 'Transcribe individual files on-add instead of batch-only'),
    ('Bulk Summarization', 'Batch-summarize completed transcriptions with Apple Intelligence'),
    ('Smart File Renaming', 'Auto-generate file names from recording date, content description, and speaker names'),
]
for name, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name} — ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'Each feature plan has been revised based on competitive analysis, Apple Human Interface Guidelines, '
    'Liquid Glass design patterns, and a thorough codebase audit.'
)

# =====================================================================
# 2. PRODUCT AUDIT
# =====================================================================
doc.add_page_break()
doc.add_heading('2. Product Audit: Current State Assessment', level=1)

doc.add_heading('2.1 Architecture Overview', level=2)

add_table(
    ['Layer', 'Technology', 'Status'],
    [
        ['UI Framework', 'SwiftUI (macOS 26+)', 'Functional'],
        ['Speech Engine', 'SpeechTranscriber / SpeechAnalyzer', 'Functional'],
        ['AI Engine', 'Apple Intelligence (Foundation Models)', 'Partially wired'],
        ['Storage', 'iCloud Drive + SQLite', 'Functional'],
        ['Concurrency', 'Swift Actors + @MainActor', 'Functional'],
    ]
)

doc.add_heading('2.2 Critical Issues (Data Loss / Broken Functionality)', level=2)
doc.add_paragraph(
    'The codebase audit uncovered 9 critical issues where functionality is broken or data may be lost:'
)

add_table(
    ['#', 'Issue', 'Severity'],
    [
        ['C1', 'Description edits not persisted to database — lost on restart', 'Critical'],
        ['C2', 'autoStartTranscription setting stored but never read', 'Critical'],
        ['C3', 'defaultLocale setting stored but never applied', 'Critical'],
        ['C4', 'defaultExportFormat setting never referenced during export', 'High'],
        ['C5', 'PersistedJobData.language stores URL instead of locale', 'Critical'],
        ['C6', 'Security-scoped resource revoked after arbitrary 3s delay', 'Critical'],
        ['C7', 'Tags/topics AI code exists but never called', 'High'],
        ['C8', 'Speaker detection AI code exists but never called', 'High'],
        ['C9', 'AI filename suggestion fully implemented but never called', 'High'],
    ]
)

doc.add_heading('2.3 High-Priority UX Issues', level=2)

add_table(
    ['#', 'Issue'],
    [
        ['U1', 'Blank right pane when processing/pending job is selected'],
        ['U2', 'No confirmation dialog for destructive Clear All / Remove actions'],
        ['U3', 'Inconsistent file addition flows (company picker only on one path)'],
        ['U4', 'Family member delete does not refresh UI'],
        ['U5', 'Chat stop button does not actually cancel AI generation'],
        ['U6', 'Chat auto-save is conditional with no visual indicator'],
        ['U7', 'No onboarding or first-run experience'],
        ['U8', 'Start button disabled with no explanation when transcription unavailable'],
    ]
)

doc.add_heading('2.4 Polish & Interaction Gaps', level=2)

add_table(
    ['#', 'Issue'],
    [
        ['P1', 'Zero keyboard shortcuts for any action'],
        ['P2', 'No context menu on transcript text'],
        ['P3', 'Three different rename entry points, inconsistent behavior'],
        ['P4', 'No animations or transitions between states'],
        ['P5', 'Inconsistent corner radii (6, 8, 12, 16 across views)'],
        ['P6', 'Hard-coded colors with no shared theme/palette'],
        ['P7', 'No drag-and-drop in Advanced mode when files exist'],
        ['P8', 'Duplicate SQLite connections in two services'],
        ['P9', 'Excessive debug print statements throughout codebase'],
    ]
)

# =====================================================================
# 3. COMPETITIVE LANDSCAPE
# =====================================================================
doc.add_page_break()
doc.add_heading('3. Competitive Landscape Analysis', level=1)

doc.add_heading('3.1 Feature Comparison Matrix', level=2)

add_table(
    ['Capability', 'NitNab', 'Otter.ai', 'MacWhisper', 'Descript', 'Rev.com'],
    [
        ['Per-file transcription', 'No', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Batch transcription', 'Yes', 'Yes', 'Yes (Pro)', 'Yes', 'Yes'],
        ['Auto-summarization', 'Partial', 'Yes (auto)', 'Via AI', 'No', 'Yes'],
        ['Bulk summarization', 'No', 'No', 'No', 'No', 'Yes'],
        ['Speaker identification', 'Not wired', 'Weak', 'Weak', 'Good', 'Good'],
        ['Smart file naming', 'Not wired', 'Calendar title', 'Manual', 'Project', 'Manual'],
        ['Offline processing', 'Yes', 'No', 'Yes', 'No', 'No'],
        ['Privacy-first', 'Yes', 'No', 'Yes', 'No', 'No'],
        ['macOS native', 'Yes', 'No', 'Yes', 'No', 'No'],
    ]
)

doc.add_heading('3.2 Key Competitive Insights', level=2)

competitors = [
    ('Otter.ai', [
        'Summaries auto-generate after meetings — users never need to request them',
        'AI Chat spans the entire transcript library, not just individual files',
        'Speaker naming remains poor ("Speaker 1", "Speaker 2") — opportunity for NitNab',
        'File naming is automatic based on meeting calendar event title',
    ]),
    ('MacWhisper', [
        'Batch transcription with simultaneous multi-format export selection',
        'Drag-and-drop opens a dedicated batch interface',
        'Described as "perfectly tailored for macOS, feeling like a native utility"',
        'Watch Folders: auto-transcribe any file placed in a designated folder',
        'Already adopted Liquid Glass design on macOS Tahoe',
    ]),
    ('Descript', [
        'Text-based editing paradigm — edit media by editing transcript text',
        'Strong multi-format workflow integrations (Adobe, Final Cut, Logic)',
        'Pricing model penalizes multi-file workflows — NitNab\'s offline model avoids this',
    ]),
    ('Rev.com', [
        'Multi-file AI chat for cross-file analysis and summarization',
        'Speaker label editing with merge functionality for duplicate speakers',
        'Legal-grade security and encryption',
        '96% AI accuracy, 99% with human verification',
    ]),
]

for name, points in competitors:
    doc.add_heading(name, level=3)
    for point in points:
        doc.add_paragraph(point, style='List Bullet')

doc.add_heading('3.3 NitNab\'s Competitive Advantages', level=2)

advantages = [
    'Fully on-device — no cloud dependency, no per-minute pricing, complete privacy',
    'Apple Intelligence integration — native Foundation Models for summarization',
    'Company context system — vocabulary and people databases for accuracy',
    'macOS native — true SwiftUI app, not Electron/web wrapper',
    'iCloud sync — structured folder output for each transcription',
]
for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_heading('3.4 Competitive Gaps to Close', level=2)

gaps = [
    'Per-file transcription is table stakes — every competitor offers it',
    'Auto-summarization should be automatic, not manual — follow Otter.ai\'s lead',
    'Speaker identification must be wired up — the code exists but is disconnected',
    'Smart naming combining date + AI description + speakers — no competitor does this well',
]
for i, gap in enumerate(gaps, 1):
    doc.add_paragraph(f'{i}. {gap}')

# =====================================================================
# 4. PLATFORM ALIGNMENT
# =====================================================================
doc.add_page_break()
doc.add_heading('4. Platform Alignment: macOS 26 Tahoe & Liquid Glass', level=1)

doc.add_heading('4.1 Liquid Glass Design Language', level=2)
doc.add_paragraph(
    'Liquid Glass is Apple\'s most significant design evolution since iOS 7, announced at WWDC 2025. '
    'It is a translucent material featuring real-time light bending (lensing), specular highlights '
    'responding to device motion, adaptive shadows, and interactive behaviors. It replaces traditional '
    'blur with optical refraction for depth and hierarchy.'
)

doc.add_paragraph(
    'For NitNab, recompiling with Xcode 26 gives standard controls Liquid Glass automatically. '
    'However, the app uses many hard-coded colors and custom backgrounds that will need to be replaced '
    'with system materials for proper Liquid Glass integration.'
)

add_table(
    ['Element', 'Current State', 'Liquid Glass Action'],
    [
        ['Toolbar', 'Standard AppKit', 'Recompile — auto-adopts'],
        ['Sidebar', 'Custom VStack', 'Adopt NavigationSplitView'],
        ['File list', 'Hard-coded colors', 'Use system list styles'],
        ['Drop zone', 'Custom rounded rect', 'Use system materials'],
        ['Settings', 'TabView', 'Auto-adopts glass tabs'],
    ]
)

doc.add_heading('4.2 Key Liquid Glass SwiftUI APIs', level=2)

add_table(
    ['API', 'Purpose'],
    [
        ['.glassEffect()', 'Apply glass material to any custom view'],
        ['.glassEffect(.regular.interactive())', 'Interactive glass responding to input'],
        ['GlassEffectContainer', 'Group glass elements; enables morphing transitions'],
        ['.glassEffectID(_:in:)', 'Associate elements for morphing across states'],
        ['.buttonStyle(.glassProminent)', 'Prominent tinted glass button style'],
        ['ToolbarSpacer', 'Control toolbar group separation'],
        ['.backgroundExtensionEffect()', 'Extend content behind floating sidebar'],
    ]
)

doc.add_heading('4.3 Foundation Models Framework', level=2)
doc.add_paragraph(
    'The Foundation Models framework (WWDC25) provides a ~3 billion parameter on-device language model '
    'via Apple Intelligence. NitNab already uses LanguageModelSession but has three disconnected functions '
    'that should be wired: extractTopics(), extractNames(), and suggestFileName(). These are critical '
    'enablers for the Bulk Summarization and Smart Renaming features.'
)

# =====================================================================
# 5. FEATURE PLAN
# =====================================================================
doc.add_page_break()
doc.add_heading('5. Feature Plan: Three Pillars of Improvement', level=1)

# -- Feature 1 --
doc.add_heading('5.1 Feature 1: Per-File Transcription', level=2)

doc.add_paragraph(
    'Competitive Context: Every competitor offers per-file transcription. NitNab\'s batch-only model '
    'is a significant friction point. MacWhisper\'s pattern — drag a file, transcription begins — is '
    'the standard macOS expectation.'
)

doc.add_heading('Implementation Components', level=3)

components_f1 = [
    ('A. ViewModel — FIFO Queue', 'Add transcribeSingleJob() method with a sequential processing queue. '
     'Per-file triggers add to front of queue. "Transcribe All" adds all pending to queue.'),
    ('B. File Row Button', 'Inline "Transcribe" button (waveform.badge.plus) per row for pending jobs. '
     'Animates to progress bar during processing.'),
    ('C. Pending State View', 'New PendingJobView shown in right pane when pending file is selected: '
     'file metadata, locale picker, company picker, and prominent "Start Transcription" CTA.'),
    ('D. Auto-Start Setting', 'Wire the existing broken autoStartTranscription setting. When enabled, '
     'automatically transcribe files as they\'re added — MacWhisper-style "drop and go."'),
    ('E. Header Update', 'Rename "Start" to "Transcribe All (N)" with count badge. Show which job '
     'is currently processing in a subtitle.'),
]
for title, desc in components_f1:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Revisions Based on Research', level=3)
revisions_f1 = [
    'Added auto-start setting wiring (discovered broken during audit)',
    'Added pending state view for right pane (blank pane was a UX gap)',
    'Added FIFO queue model (inspired by MacWhisper\'s sequential batch)',
    'Per-file button placed both in row AND in detail pane for discoverability',
]
for r in revisions_f1:
    doc.add_paragraph(r, style='List Bullet')

# -- Feature 2 --
doc.add_heading('5.2 Feature 2: Bulk Summarization', level=2)

doc.add_paragraph(
    'Competitive Context: Otter.ai auto-generates summaries after every meeting. Rev.com allows '
    'multi-file AI chat. No competitor offers explicit "bulk summarize" — this is a differentiation opportunity.'
)

doc.add_heading('Implementation Components', level=3)

components_f2 = [
    ('A. Data Model', 'Add summary, summaryGeneratedAt, and summaryStatus fields to TranscriptionJob. '
     'Database migration for new columns.'),
    ('B. ViewModel Methods', 'summarizeSingleJob(), summarizeAllCompleted(), and autoSummarizeAfterTranscription(). '
     'Also wire extractTopics() and extractNames() alongside each summarization.'),
    ('C. Bulk Button', '"Summarize All (N)" button in HeaderView. Shows progress during operation.'),
    ('D. Summary Display', 'Summary card displayed ABOVE transcript tabs (not buried in a tab). '
     'Collapsible with disclosure triangle. Includes "Regenerate" option.'),
    ('E. File Row Indicators', 'Sparkles icon for summarized files. Tooltip: "AI Summary available."'),
    ('F. Settings', 'Auto-summarize toggle and auto-extract topics/speakers toggle.'),
]
for title, desc in components_f2:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Revisions Based on Research', level=3)
revisions_f2 = [
    'Auto-summarize after transcription (inspired by Otter.ai\'s automatic approach)',
    'Wire up extractTopics() and extractNames() alongside summarization (found disconnected in audit)',
    'Summary placement moved above tabs (Otter shows summary at top of notes)',
    'Tag cloud population enabled (audit found tags were always empty)',
]
for r in revisions_f2:
    doc.add_paragraph(r, style='List Bullet')

# -- Feature 3 --
doc.add_heading('5.3 Feature 3: Smart File Renaming', level=2)

doc.add_paragraph(
    'Competitive Context: No competitor does smart naming well. Otter uses calendar titles. '
    'MacWhisper and Rev use original filenames. NitNab can differentiate by offering '
    'AI-powered multi-component naming with user-configurable formula.'
)

doc.add_heading('Implementation Components', level=3)

components_f3 = [
    ('A. NamingPreferences Model', 'User-configurable struct: includeDate, includeDescription, '
     'includeSpeakerNames, dateFormat, descriptionStyle, nameStyle, separator, maxLength, autoRename.'),
    ('B. AI Name Generation', 'Enhance AIService.suggestFileName() to accept NamingPreferences. '
     'Assemble components: "2026-02-15 — Q1 Budget Review — Sarah, Mike"'),
    ('C. Smart Rename Sheet', 'Sheet with live preview, component toggles, editable name field, '
     'Apply/Skip buttons, and "Always use these settings" option.'),
    ('D. Post-Transcription Flow', 'If auto-rename ON: apply silently with notification. '
     'If OFF: show SmartRenameSheet with suggestion.'),
    ('E. Bulk Rename', 'bulkSmartRename() with review sheet showing all proposed renames before applying.'),
    ('F. Settings', 'File Naming section with component toggles, format pickers, separator picker, '
     'and live preview showing example output.'),
]
for title, desc in components_f3:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Revisions Based on Research', level=3)
revisions_f3 = [
    'Live preview in settings and rename sheet (no competitor shows this)',
    'Bulk rename with review sheet (never rename silently in bulk)',
    'Auto-rename defaults to OFF (user consent paramount per HIG)',
    'Name style options (first/full/initials) based on Rev.com\'s speaker label flexibility',
]
for r in revisions_f3:
    doc.add_paragraph(r, style='List Bullet')

# =====================================================================
# 6. UX POLISH
# =====================================================================
doc.add_page_break()
doc.add_heading('6. UX Polish & Usability Recommendations', level=1)

doc.add_heading('6.1 Keyboard Shortcuts', level=2)
doc.add_paragraph('The app currently has zero keyboard shortcuts. Every professional macOS app provides these:')

add_table(
    ['Action', 'Shortcut', 'Rationale'],
    [
        ['Add files', 'Cmd+O', 'Standard macOS file open'],
        ['Transcribe selected', 'Cmd+Return', 'Natural "go" action'],
        ['Transcribe all', 'Cmd+Shift+Return', 'Batch variant'],
        ['Copy transcript', 'Cmd+Shift+C', 'Transcript context'],
        ['Export', 'Cmd+E', 'Standard export'],
        ['Search', 'Cmd+F', 'Standard search'],
        ['Toggle mode', 'Cmd+1 / Cmd+2', 'Mode switching'],
        ['Delete selected', 'Cmd+Delete', 'With confirmation'],
        ['Rename', 'Return', 'Finder convention'],
        ['Settings', 'Cmd+,', 'Standard macOS'],
    ]
)

doc.add_heading('6.2 Eliminate Blank Panes', level=2)
doc.add_paragraph('Every job status should have a corresponding right-pane view:')

add_table(
    ['Job Status', 'Right Pane Content'],
    [
        ['No selection', 'Welcome/drop zone with instructions'],
        ['Pending', 'File metadata + "Start Transcription" CTA'],
        ['Processing', 'Live progress bar + elapsed time + cancel'],
        ['Completed', 'Transcript view (existing)'],
        ['Failed', 'Error details + Retry button'],
    ]
)

doc.add_heading('6.3 Confirmation Dialogs', level=2)
doc.add_paragraph(
    'All destructive actions must require confirmation. Currently, Remove, Clear Completed, '
    'and Clear All proceed silently with no undo.'
)

doc.add_heading('6.4 Design System', level=2)
doc.add_paragraph(
    'Create a shared DesignSystem constants file to replace hard-coded corner radii (6, 8, 12, 16) '
    'and colors (Color.blue.opacity(0.05), etc.) with consistent values. Adopt system materials '
    'for Liquid Glass compatibility.'
)

# =====================================================================
# 7. ACCESSIBILITY
# =====================================================================
doc.add_page_break()
doc.add_heading('7. Accessibility & Inclusivity', level=1)

add_table(
    ['Area', 'Current State', 'Recommendation'],
    [
        ['VoiceOver labels', 'Basic', 'Add descriptive labels with status context'],
        ['Progress announcements', 'None', 'Use AccessibilityNotification.Announcement'],
        ['Keyboard navigation', 'Not configured', 'Full tab/arrow key support'],
        ['High Contrast', 'Not tested', 'Test all views with Increase Contrast'],
        ['Reduce Transparency', 'Not tested', 'Verify Liquid Glass readability'],
        ['Dynamic Type', 'Not supported', 'Support for any custom text sizing'],
    ]
)

# =====================================================================
# 8. IMPLEMENTATION ROADMAP
# =====================================================================
doc.add_page_break()
doc.add_heading('8. Implementation Roadmap', level=1)

phases = [
    ('Phase 1: Foundation & Bug Fixes (Week 1)', [
        ['Fix critical bugs C1-C6', 'Critical', '8 hours'],
        ['Add confirmation dialogs', 'High', '2 hours'],
        ['Unify file addition flows', 'High', '3 hours'],
        ['Create DesignSystem constants', 'Medium', '2 hours'],
        ['Remove debug artifacts', 'Medium', '1 hour'],
    ]),
    ('Phase 2: Per-File Transcription (Week 2)', [
        ['FIFO transcription queue', 'High', '4 hours'],
        ['transcribeSingleJob() method', 'High', '2 hours'],
        ['Per-row Transcribe button', 'High', '2 hours'],
        ['PendingJobView + ProcessingJobView', 'High', '5 hours'],
        ['Header update + keyboard shortcuts', 'Medium', '3 hours'],
    ]),
    ('Phase 3: Bulk Summarization (Week 3)', [
        ['Summary data model + DB migration', 'High', '3 hours'],
        ['Summarize methods + wire AI functions', 'High', '7 hours'],
        ['Summary card UI above tabs', 'High', '3 hours'],
        ['Summarize All button + indicators', 'Medium', '2 hours'],
        ['Auto-summarize setting', 'Medium', '1 hour'],
    ]),
    ('Phase 4: Smart File Renaming (Week 4)', [
        ['NamingPreferences model', 'High', '1 hour'],
        ['Enhanced suggestFileName()', 'High', '3 hours'],
        ['SmartRenameSheet view', 'High', '4 hours'],
        ['Settings + post-transcription flow', 'High', '5 hours'],
        ['Bulk rename + context menu', 'Medium', '4 hours'],
    ]),
    ('Phase 5: Polish & Accessibility (Week 5)', [
        ['Liquid Glass adoption (NavigationSplitView, materials)', 'High', '7 hours'],
        ['Keyboard shortcuts (full set)', 'High', '3 hours'],
        ['VoiceOver audit + labels', 'High', '4 hours'],
        ['Animations + transitions', 'Medium', '3 hours'],
        ['Onboarding flow', 'Medium', '4 hours'],
    ]),
]

for phase_title, tasks in phases:
    doc.add_heading(phase_title, level=2)
    add_table(
        ['Task', 'Priority', 'Effort'],
        tasks
    )

# =====================================================================
# 9. APPENDIX
# =====================================================================
doc.add_page_break()
doc.add_heading('9. Appendix: Sources', level=1)

sources = [
    'Apple Human Interface Guidelines — developer.apple.com/design/human-interface-guidelines/',
    'Apple Liquid Glass Announcement — apple.com/newsroom/2025/06/',
    'Foundation Models Framework WWDC25 — developer.apple.com/videos/play/wwdc2025/286/',
    'Liquid Glass in Swift Best Practices — dev.to/diskcleankit/',
    'SpeechAnalyzer WWDC25 — developer.apple.com/videos/play/wwdc2025/277/',
    'Otter.ai Review (2026) — tldv.io/blog/otter-ai-review/',
    'MacWhisper Batch Transcription — macwhisper.helpscoutdocs.com/',
    'Descript Review (2025) — castmagic.io/software-review/descript',
    'Rev.com Transcription Editor — support.rev.com/',
]

for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# =====================================================================
# SAVE
# =====================================================================
doc.save('/home/user/nitnab/docs/NitNab_Design_Review_and_Feature_Plan.docx')
print('DOCX generated successfully.')
